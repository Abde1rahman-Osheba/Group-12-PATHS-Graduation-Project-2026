"""
PATHS Backend — CV Ingestion Agent Nodes (LangGraph).

Each function is a node in the LangGraph workflow. All nodes receive the
CVIngestionState dict and return a partial state update dict.
"""

import os
import re
import json
import uuid
import hashlib
import logging
from typing import Dict, Any, List

from langchain_text_splitters import RecursiveCharacterTextSplitter

from pypdf import PdfReader
from app.agents.cv_ingestion.state import CVIngestionState
from app.agents.cv_ingestion.schemas import CandidateExtraction
from app.core.config import get_settings
from app.core.database import SessionLocal
from app.db.models.candidate import Candidate
from app.db.models.cv_entities import (
    CandidateDocument, Skill, CandidateSkill, CandidateExperience,
    CandidateEducation, CandidateCertification
)
from app.db.models.evidence import EvidenceItem, CandidateSource
from app.repositories import graph_repo, vector_repo
from app.services.embedding_service import embed_documents
from app.services.llm.openrouter_client import generate_json_response

logger = logging.getLogger(__name__)
settings = get_settings()

# ────────────────────────────────────────────────────
# 1. load_document
# ────────────────────────────────────────────────────

def load_document(state: CVIngestionState) -> Dict:
    """Load and extract raw text from the uploaded document."""
    file_path = state["file_path"]
    job_id = state["job_id"]
    logger.info("[job=%s] Loading document: %s", job_id, file_path)
    try:
        raw_text = ""
        if file_path.lower().endswith(".pdf"):
            reader = PdfReader(file_path)
            for page in reader.pages:
                raw_text += (page.extract_text() or "") + "\n"
        elif file_path.lower().endswith(".docx"):
            import docx
            doc = docx.Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()

        if not raw_text.strip():
            logger.error("[job=%s] No text extracted from document", job_id)
            return {"errors": state["errors"] + ["No text could be extracted from document."], "status": "failed"}

        logger.info("[job=%s] Extracted %d chars of text", job_id, len(raw_text))
        return {"raw_text": raw_text, "stage": "extracted_text"}
    except Exception as e:
        logger.exception("[job=%s] Load document error", job_id)
        return {"errors": state["errors"] + [f"Load document error: {str(e)}"], "status": "failed"}


# ────────────────────────────────────────────────────
# 2. extract_structured_candidate_data
# ────────────────────────────────────────────────────

# Curated technical-skill keywords for offline, deterministic skill detection.
# This is the reliable baseline: it needs no LLM and no network, so a CV's
# skills are extracted even when the cloud LLM is unreachable.
_SKILL_KEYWORDS: List[str] = [
    "python", "java", "javascript", "typescript", "golang", "rust", "ruby",
    "php", "kotlin", "swift", "scala", "c++", "c#", ".net", "sql", "bash",
    "react", "angular", "vue", "next.js", "node.js", "express", "django",
    "flask", "fastapi", "spring boot", "spring", "laravel", "rails",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
    "opencv", "machine learning", "deep learning", "computer vision", "nlp",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite",
    "oracle", "cassandra", "neo4j", "qdrant",
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "terraform",
    "ansible", "jenkins", "git", "github", "gitlab", "ci/cd", "linux",
    "html", "css", "tailwind", "sass", "bootstrap", "graphql", "rest api",
    "kafka", "rabbitmq", "spark", "hadoop", "airflow",
    "agile", "scrum", "jira", "figma", "power bi", "tableau",
]


def _match_skills(raw_text: str) -> List[dict]:
    """Detect known skills in CV text by keyword — fully offline, no LLM."""
    low = (raw_text or "").lower()
    found: List[dict] = []
    seen: set = set()
    for kw in _SKILL_KEYWORDS:
        if kw in seen:
            continue
        # Boundary-aware so a short token does not match inside another word
        # (e.g. "java" must not match inside "javascript").
        pattern = r'(?<![a-z0-9+#.])' + re.escape(kw) + r'(?![a-z0-9+#])'
        if re.search(pattern, low):
            found.append({"name": kw, "category": "technical"})
            seen.add(kw)
    return found


def _deterministic_extract(raw_text: str) -> dict:
    """Extract obvious fields with no LLM: emails, phones, URLs, known skills."""
    result: dict = {}
    text = raw_text or ""

    # Email
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if emails:
        result["email"] = emails[0]

    # Phone
    phones = re.findall(r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\./0-9]{7,15}', text)
    if phones:
        result["phone"] = phones[0].strip()

    # URLs (LinkedIn, GitHub, portfolio)
    result["links"] = re.findall(r'https?://[^\s<>"\']+', text)

    # Known-skill keyword detection — the offline baseline.
    result["skills"] = _match_skills(text)

    return result


def extract_structured_candidate_data(state: CVIngestionState) -> Dict:
    """Extract structured candidate data via OpenRouter (cloud LLM) + regex.

    Uses OpenRouter rather than the local Ollama model. Structured extraction
    over a whole CV on a CPU-bound 8B model has no bounded runtime — the call
    never returned, so it hung the entire ingestion pipeline indefinitely and
    the failure handler never ran. OpenRouter is fast and its HTTP call is
    bounded by a timeout; on any failure we still degrade to deterministic
    extraction so the upload always completes and the CV is saved.
    """
    job_id = state["job_id"]
    logger.info("[job=%s] Extracting structured candidate data", job_id)
    raw_text = state.get("raw_text") or ""
    # Regex extraction of unambiguous fields — runs regardless of the LLM.
    deterministic = _deterministic_extract(raw_text)
    try:
        system_prompt = (
            "You are an expert HR CV parser. Extract ALL candidate information "
            "from the CV text and return ONLY a single JSON object — no markdown "
            "and no commentary. Use exactly these keys: full_name (string), "
            "email (string or null), phone (string or null), location_text "
            "(string or null), summary (string or null), years_experience "
            "(integer or null), skills (array of objects {name, category}), "
            "experiences (array of objects {company_name, title, start_date, "
            "end_date, description}), education (array of objects {institution, "
            "degree, field_of_study, start_date, end_date}), certifications "
            "(array of objects {name, issuer}). Be thorough: extract every "
            "skill, role, degree and certification. Use null for unknown fields."
        )
        # Cap the CV text so the request stays well within the model context.
        user_prompt = f"CV text:\n\n{raw_text[:12000]}"
        raw = generate_json_response(
            system_prompt,
            user_prompt,
            model=settings.openrouter_model,
            temperature=0.0,
            # Generous cap so a CV with many roles/skills is not truncated
            # mid-JSON (which would silently drop the trailing experiences).
            max_tokens=8000,
        )
        # Validate against the schema; if the model returned a slightly off
        # shape keep the raw dict — normalize_entities is fully defensive.
        try:
            structured = CandidateExtraction.model_validate(raw).model_dump()
        except Exception:  # noqa: BLE001
            logger.warning("[job=%s] Extraction JSON off-schema — using raw dict", job_id)
            structured = raw if isinstance(raw, dict) else {}

        # Deterministic results win for email / phone (regex is precise).
        if deterministic.get("email"):
            structured["email"] = deterministic["email"]
        if deterministic.get("phone"):
            structured["phone"] = deterministic["phone"]

        # If the LLM returned no skills (or never ran), use the offline
        # keyword-matched skills so the skills list is always populated.
        if not structured.get("skills"):
            structured["skills"] = deterministic.get("skills") or []

        logger.info(
            "[job=%s] Extraction complete: name=%s, %d skills, %d experiences, %d education",
            job_id,
            structured.get("full_name"),
            len(structured.get("skills") or []),
            len(structured.get("experiences") or []),
            len(structured.get("education") or []),
        )
        return {"structured_candidate": structured, "stage": "extracted_structure"}

    except Exception as e:  # noqa: BLE001
        # Never fail the whole upload — persist the CV with whatever the regex
        # pass found. full_name is left blank so persist_postgres keeps the
        # candidate's existing name (it only overwrites with a truthy value).
        logger.warning(
            "[job=%s] LLM extraction unavailable (%s) — using offline keyword extraction",
            job_id, e,
        )
        return {
            "structured_candidate": {
                "full_name": "",
                "email": deterministic.get("email"),
                "phone": deterministic.get("phone"),
                "location_text": None,
                "summary": None,
                "years_experience": None,
                "skills": deterministic.get("skills") or [],
                "experiences": [],
                "education": [],
                "certifications": [],
            },
            "stage": "extracted_structure",
            "errors": state["errors"] + [f"LLM extraction failed, used fallback: {str(e)}"],
        }


# ────────────────────────────────────────────────────
# 3. normalize_entities
# ────────────────────────────────────────────────────

def normalize_entities(state: CVIngestionState) -> Dict:
    """Normalize and deduplicate extracted entities."""
    job_id = state["job_id"]
    structured = state.get("structured_candidate")
    if not structured:
        return {"errors": state["errors"] + ["No structured data to normalize."], "status": "failed"}

    logger.info("[job=%s] Normalizing entities", job_id)

    # Normalize skills — accept both {"name": ...} objects and bare strings,
    # then lowercase / trim / dedupe.
    normalized_skills = []
    seen = set()
    for s in structured.get("skills", []):
        if isinstance(s, str):
            n_name, category = s.strip().lower(), None
        elif isinstance(s, dict):
            n_name = (s.get("name") or s.get("skill") or "").strip().lower()
            category = (s.get("category") or "").strip() or None
        else:
            continue
        if n_name and n_name not in seen:
            seen.add(n_name)
            normalized_skills.append({"name": n_name, "category": category})
    structured["skills"] = normalized_skills

    # Clean experiences — tolerate the field-name variations different LLMs
    # use, and keep an entry that has at least a title OR a company (interns,
    # freelancers and project roles often have only one of the two). Dropping
    # on "company AND title" silently lost real experiences.
    clean_exps = []
    for exp in structured.get("experiences", []):
        if not isinstance(exp, dict):
            continue
        company = (
            exp.get("company_name") or exp.get("company")
            or exp.get("organization") or exp.get("employer") or ""
        ).strip()
        title = (
            exp.get("title") or exp.get("role")
            or exp.get("position") or exp.get("job_title") or ""
        ).strip()
        if not (company or title):
            continue
        clean_exps.append({
            "company_name": company or "—",
            "title": title or "—",
            "start_date": (exp.get("start_date") or exp.get("from") or "").strip() or None,
            "end_date": (exp.get("end_date") or exp.get("to") or "").strip() or None,
            "description": (exp.get("description") or exp.get("summary") or "").strip() or None,
        })
    structured["experiences"] = clean_exps

    # Clean education — same tolerance for field-name variations.
    clean_edu = []
    for edu in structured.get("education", []):
        if not isinstance(edu, dict):
            continue
        institution = (
            edu.get("institution") or edu.get("school") or edu.get("university") or ""
        ).strip()
        degree = (edu.get("degree") or edu.get("qualification") or "").strip()
        if not (institution or degree):
            continue
        clean_edu.append({
            "institution": institution or "—",
            "degree": degree or None,
            "field_of_study": (
                edu.get("field_of_study") or edu.get("field") or edu.get("major") or ""
            ).strip() or None,
            "start_date": (edu.get("start_date") or "").strip() or None,
            "end_date": (edu.get("end_date") or edu.get("graduation_year") or "").strip() or None,
        })
    structured["education"] = clean_edu

    # Clean certifications
    clean_certs = []
    for cert in structured.get("certifications", []):
        name = (cert.get("name") or "").strip()
        if name:
            clean_certs.append({
                "name": name,
                "issuer": (cert.get("issuer") or "").strip() or None,
            })
    structured["certifications"] = clean_certs

    # Clean core fields
    structured["full_name"] = (structured.get("full_name") or "Unknown").strip()
    structured["email"] = (structured.get("email") or "").strip() or None
    structured["phone"] = (structured.get("phone") or "").strip() or None
    structured["location_text"] = (structured.get("location_text") or "").strip() or None
    structured["summary"] = (structured.get("summary") or "").strip() or None

    logger.info("[job=%s] Normalized: %d skills, %d exps, %d edu, %d certs",
                job_id, len(normalized_skills), len(clean_exps), len(clean_edu), len(clean_certs))

    return {"normalized_candidate": structured, "stage": "normalized_entities"}


# ────────────────────────────────────────────────────
# 4. persist_postgres
# ────────────────────────────────────────────────────

def persist_postgres(state: CVIngestionState) -> Dict:
    """Persist all normalized candidate data to PostgreSQL in one transaction."""
    job_id = state["job_id"]
    normalized = state.get("normalized_candidate")
    if not normalized:
        return {"errors": state["errors"] + ["No normalized data to persist."], "status": "failed"}

    candidate_id = state.get("candidate_id") or str(uuid.uuid4())
    doc_id = state.get("document_id") or str(uuid.uuid4())

    logger.info("[job=%s][candidate=%s] Persisting to PostgreSQL", job_id, candidate_id)

    db = SessionLocal()
    try:
        candidate_uuid = uuid.UUID(candidate_id)
        from sqlalchemy import select

        # ── Candidate core (upsert) ──
        candidate = db.get(Candidate, candidate_uuid)
        if not candidate:
            candidate = Candidate(
                id=candidate_uuid,
                full_name=normalized.get("full_name") or "Unknown",
                email=normalized.get("email"),
                phone=normalized.get("phone"),
                location_text=normalized.get("location_text"),
                summary=normalized.get("summary"),
                years_experience=normalized.get("years_experience"),
                status="active",
            )
            db.add(candidate)
        else:
            candidate.full_name = normalized.get("full_name") or candidate.full_name
            candidate.email = normalized.get("email") or candidate.email
            candidate.phone = normalized.get("phone") or candidate.phone
            candidate.location_text = normalized.get("location_text") or candidate.location_text
            candidate.summary = normalized.get("summary") or candidate.summary
            candidate.years_experience = normalized.get("years_experience") or candidate.years_experience
        db.flush()

        # Mirror extracted skill names onto Candidate.skills — the array column
        # the candidate portal profile (GET /me/profile) actually reads. Without
        # this the CV's skills would only live in the normalized candidate_skills
        # table and never appear on the candidate's profile page. Existing skills
        # are kept and de-duplicated (case-insensitive).
        extracted_skill_names = [
            (s.get("name") or "").strip()
            for s in normalized.get("skills", [])
            if (s.get("name") or "").strip()
        ]
        if extracted_skill_names:
            merged = list(candidate.skills or [])
            seen = {x.lower() for x in merged}
            for nm in extracted_skill_names:
                if nm.lower() not in seen:
                    merged.append(nm)
                    seen.add(nm.lower())
            candidate.skills = merged[:100]

        # ── Candidate document ──
        raw_text = state.get("raw_text", "")
        checksum = hashlib.sha256(raw_text.encode("utf-8")).hexdigest() if raw_text else None

        # Determine mime type
        file_path = state.get("file_path", "")
        ext = os.path.splitext(file_path)[1].lower()
        mime_map = {".pdf": "application/pdf", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".txt": "text/plain"}
        mime_type = mime_map.get(ext, "application/octet-stream")

        # The upload endpoint stores files as "<uuid>_<original-name>" — strip
        # that prefix so the candidate sees the real name they uploaded.
        stored_name = os.path.basename(file_path)
        clean_name = re.sub(
            r"^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-"
            r"[0-9a-fA-F]{4}-[0-9a-fA-F]{12}_",
            "",
            stored_name,
        )
        # Upsert: the upload endpoint may have already created this document
        # row (so the file shows in "Uploaded Files" immediately, before
        # extraction). Enrich it here instead of inserting a duplicate PK.
        existing_doc = db.get(CandidateDocument, uuid.UUID(doc_id))
        if existing_doc is not None:
            existing_doc.raw_text = raw_text
            existing_doc.checksum = checksum
            if not existing_doc.original_filename:
                existing_doc.original_filename = clean_name or stored_name
        else:
            db.add(CandidateDocument(
                id=uuid.UUID(doc_id),
                candidate_id=candidate_uuid,
                original_filename=clean_name or stored_name,
                mime_type=mime_type,
                storage_path_or_url=file_path,
                raw_text=raw_text,
                checksum=checksum,
            ))

        # ── Skills (upsert skill catalog + link) ──
        for skill_data in normalized.get("skills", []):
            skill_name = skill_data.get("name")
            if not skill_name:
                continue

            # Use .first() (not scalar_one_or_none) — the skills catalogue can
            # contain duplicate normalized_name rows when the unique constraint
            # wasn't enforced, and scalar_one_or_none() RAISES on >1 row, which
            # previously aborted the whole ingestion ("Multiple rows were found").
            existing_skill = db.execute(
                select(Skill).where(Skill.normalized_name == skill_name)
            ).scalars().first()
            if not existing_skill:
                existing_skill = Skill(normalized_name=skill_name, category=skill_data.get("category"))
                db.add(existing_skill)
                db.flush()

            # Check if link already exists (duplicate-safe for the same reason)
            existing_link = db.execute(
                select(CandidateSkill).where(
                    CandidateSkill.candidate_id == candidate_uuid,
                    CandidateSkill.skill_id == existing_skill.id,
                )
            ).scalars().first()
            if not existing_link:
                cs = CandidateSkill(candidate_id=candidate_uuid, skill_id=existing_skill.id)
                db.add(cs)

        # ── Experiences (skip rows already on the candidate) ──
        existing_exps = db.execute(
            select(CandidateExperience.company_name, CandidateExperience.title)
            .where(CandidateExperience.candidate_id == candidate_uuid)
        ).all()
        exp_seen = {
            ((c or "").strip().lower(), (t or "").strip().lower())
            for c, t in existing_exps
        }
        for exp in normalized.get("experiences", []):
            company = exp.get("company_name") or "Unknown"
            title = exp.get("title") or "Unknown Role"
            key = (company.strip().lower(), title.strip().lower())
            if key in exp_seen:
                continue
            exp_seen.add(key)
            db.add(CandidateExperience(
                candidate_id=candidate_uuid,
                company_name=company,
                title=title,
                start_date=exp.get("start_date"),
                end_date=exp.get("end_date"),
                description=exp.get("description"),
            ))

        # ── Education (skip rows already on the candidate) ──
        existing_edu = db.execute(
            select(CandidateEducation.institution, CandidateEducation.degree)
            .where(CandidateEducation.candidate_id == candidate_uuid)
        ).all()
        edu_seen = {
            ((i or "").strip().lower(), (dg or "").strip().lower())
            for i, dg in existing_edu
        }
        for edu in normalized.get("education", []):
            institution = edu.get("institution") or "Unknown"
            degree = edu.get("degree") or ""
            key = (institution.strip().lower(), degree.strip().lower())
            if key in edu_seen:
                continue
            edu_seen.add(key)
            db.add(CandidateEducation(
                candidate_id=candidate_uuid,
                institution=institution,
                degree=edu.get("degree"),
                field_of_study=edu.get("field_of_study"),
                start_date=edu.get("start_date"),
                end_date=edu.get("end_date"),
            ))

        # ── Certifications ──
        for cert in normalized.get("certifications", []):
            db.add(CandidateCertification(
                candidate_id=candidate_uuid,
                name=cert.get("name", "Unknown"),
                issuer=cert.get("issuer"),
            ))

        # ── Evidence items (Blueprint Law #1 — evidence over inference) ──────
        # Write one evidence_item per extracted entity so every agent claim
        # can be traced back to a concrete artefact in the source CV.
        file_path_val = state.get("file_path", "")

        # One source row for the uploaded CV document itself
        db.add(CandidateSource(
            id=uuid.uuid4(),
            candidate_id=candidate_uuid,
            source="cv",
            url=None,
            raw_blob_uri=file_path_val or None,
            fetched_at=None,
        ))

        # Skill evidence items
        for skill_data in normalized.get("skills", []):
            skill_name = skill_data.get("name")
            if not skill_name:
                continue
            db.add(EvidenceItem(
                id=uuid.uuid4(),
                candidate_id=candidate_uuid,
                ingestion_job_id=job_id,
                type="cv_claim",
                field_ref=f"skill:{skill_name}",
                source_uri=file_path_val or None,
                extracted_text=skill_name,
                confidence=0.9,
                meta_json={"category": skill_data.get("category")},
            ))

        # Experience evidence items
        for idx, exp in enumerate(normalized.get("experiences", [])):
            db.add(EvidenceItem(
                id=uuid.uuid4(),
                candidate_id=candidate_uuid,
                ingestion_job_id=job_id,
                type="cv_claim",
                field_ref=f"experience:{idx}",
                source_uri=file_path_val or None,
                extracted_text=exp.get("description") or exp.get("title", ""),
                confidence=0.85,
                meta_json={
                    "title": exp.get("title"),
                    "company": exp.get("company_name"),
                    "start_date": exp.get("start_date"),
                    "end_date": exp.get("end_date"),
                },
            ))

        # Education evidence items
        for idx, edu in enumerate(normalized.get("education", [])):
            db.add(EvidenceItem(
                id=uuid.uuid4(),
                candidate_id=candidate_uuid,
                ingestion_job_id=job_id,
                type="cv_claim",
                field_ref=f"education:{idx}",
                source_uri=file_path_val or None,
                extracted_text=f"{edu.get('degree', '')} {edu.get('field_of_study', '')}".strip(),
                confidence=0.9,
                meta_json={
                    "institution": edu.get("institution"),
                    "degree": edu.get("degree"),
                    "field_of_study": edu.get("field_of_study"),
                },
            ))

        # Certification evidence items
        for cert in normalized.get("certifications", []):
            cert_name = cert.get("name")
            if not cert_name:
                continue
            db.add(EvidenceItem(
                id=uuid.uuid4(),
                candidate_id=candidate_uuid,
                ingestion_job_id=job_id,
                type="cv_claim",
                field_ref=f"certification:{cert_name}",
                source_uri=file_path_val or None,
                extracted_text=cert_name,
                confidence=0.9,
                meta_json={"issuer": cert.get("issuer")},
            ))

        db.commit()
        logger.info("[job=%s][candidate=%s] PostgreSQL persistence complete", job_id, candidate_id)
        return {"candidate_id": candidate_id, "document_id": doc_id, "stage": "persisted_postgres"}

    except Exception as e:
        db.rollback()
        logger.exception("[job=%s] Postgres persistence error", job_id)
        return {"errors": state["errors"] + [f"Postgres error: {str(e)}"], "status": "failed"}
    finally:
        db.close()


# ────────────────────────────────────────────────────
# 5. project_to_age
# ────────────────────────────────────────────────────

def project_to_age(state: CVIngestionState) -> Dict:
    """Project all candidate relationships into the Apache AGE graph."""
    job_id = state["job_id"]
    candidate_id = state.get("candidate_id")
    document_id = state.get("document_id")
    normalized = state.get("normalized_candidate")

    logger.info("[job=%s][candidate=%s] Projecting to AGE graph", job_id, candidate_id)
    db = SessionLocal()
    try:
        graph_repo.init_graph(db)

        # Candidate node
        graph_repo.project_candidate(db, candidate_id, normalized)

        # Skills
        for skill in normalized.get("skills", []):
            graph_repo.project_skill(db, candidate_id, skill["name"], skill["name"])

        # Companies / work experience
        for exp in normalized.get("experiences", []):
            company = exp.get("company_name")
            if company:
                graph_repo.project_company(db, candidate_id, company, exp.get("title", ""))

        # Education
        for edu in normalized.get("education", []):
            institution = edu.get("institution")
            if institution:
                graph_repo.project_education(db, candidate_id, institution, edu.get("degree", ""))

        # Certifications
        for cert in normalized.get("certifications", []):
            name = cert.get("name")
            if name:
                graph_repo.project_certification(db, candidate_id, name, cert.get("issuer", ""))

        # Document linkage
        if document_id:
            graph_repo.project_document(db, candidate_id, document_id)

        db.commit()
        logger.info("[job=%s] AGE projection complete", job_id)
        return {"stage": "projected_age"}
    except Exception as e:
        db.rollback()
        logger.exception("[job=%s] AGE projection error", job_id)
        # AGE failure after PG commit: mark stage failed but allow retry
        return {"errors": state["errors"] + [f"AGE error: {str(e)}"], "stage": "age_failed"}
    finally:
        db.close()


# ────────────────────────────────────────────────────
# 6. chunk_document
# ────────────────────────────────────────────────────

def chunk_document(state: CVIngestionState) -> Dict:
    """Split the raw CV text into chunks preserving section awareness."""
    job_id = state["job_id"]
    raw_text = state.get("raw_text", "")
    logger.info("[job=%s] Chunking document (%d chars)", job_id, len(raw_text))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(raw_text)
    logger.info("[job=%s] Created %d chunks", job_id, len(chunks))
    return {"chunks": chunks, "stage": "chunked"}


# ────────────────────────────────────────────────────
# 7. embed_chunks
# ────────────────────────────────────────────────────

def embed_chunks(state: CVIngestionState) -> Dict:
    """Embed all text chunks using the local Ollama embedding model."""
    job_id = state["job_id"]
    chunks = state.get("chunks", [])
    if not chunks:
        logger.warning("[job=%s] No chunks to embed", job_id)
        return {"embeddings": [], "stage": "embedded"}
    try:
        logger.info("[job=%s] Embedding %d chunks", job_id, len(chunks))
        embeddings = embed_documents(chunks)
        logger.info("[job=%s] Embedding complete, dim=%d", job_id, len(embeddings[0]) if embeddings else 0)
        return {"embeddings": embeddings, "stage": "embedded"}
    except Exception as e:
        logger.exception("[job=%s] Embedding error", job_id)
        return {"errors": state["errors"] + [f"Embedding error: {str(e)}"], "stage": "embed_failed"}


# ────────────────────────────────────────────────────
# 8. upsert_qdrant
# ────────────────────────────────────────────────────

def upsert_qdrant(state: CVIngestionState) -> Dict:
    """Upsert embedded chunks to Qdrant vector store."""
    job_id = state["job_id"]
    chunks = state.get("chunks", [])
    embeddings = state.get("embeddings", [])
    cand_id = state.get("candidate_id")
    doc_id = state.get("document_id")

    if chunks and embeddings:
        try:
            logger.info("[job=%s] Upserting %d vectors to Qdrant", job_id, len(chunks))
            vector_repo.init_collection()
            vector_repo.upsert_chunks(cand_id, doc_id, chunks, embeddings)
            logger.info("[job=%s] Qdrant upsert complete", job_id)
            return {"stage": "upserted_qdrant"}
        except Exception as e:
            logger.exception("[job=%s] Qdrant upsert error", job_id)
            return {"errors": state["errors"] + [f"Qdrant error: {str(e)}"], "stage": "qdrant_failed"}
    logger.warning("[job=%s] No chunks/embeddings to upsert", job_id)
    return {"stage": "upserted_qdrant"}


# ────────────────────────────────────────────────────
# 9a. sync_unified_candidate  (one-vector-per-candidate spec)
# ────────────────────────────────────────────────────

def sync_unified_candidate(state: CVIngestionState) -> Dict:
    """Project the candidate to AGE + Qdrant per the unified-ID spec.

    Produces exactly one Qdrant point in `QDRANT_CANDIDATE_COLLECTION`
    (`paths_candidates`) using the PostgreSQL `candidate_id` as both
    the point ID and the payload `candidate_id`. Failures here do not
    fail the pipeline — they are recorded in `db_sync_status` so the
    `/admin/sync/candidate/{id}/retry` endpoint can recover.
    """
    job_id = state["job_id"]
    candidate_id = state.get("candidate_id")
    if not candidate_id:
        return {"stage": "unified_sync_skipped"}

    db = SessionLocal()
    try:
        from app.services.candidate_sync_service import sync_candidate_full

        result = sync_candidate_full(db, candidate_id)
        logger.info(
            "[job=%s][candidate=%s] unified sync result: %s",
            job_id, candidate_id, result,
        )
        return {"stage": "unified_synced"}
    except Exception as exc:
        logger.exception(
            "[job=%s][candidate=%s] unified candidate sync failed",
            job_id, candidate_id,
        )
        # Spec rule: never lose the PostgreSQL row on sync failure.
        return {
            "errors": state["errors"] + [f"Unified sync error: {str(exc)}"],
            "stage": "unified_sync_failed",
        }
    finally:
        db.close()


# ────────────────────────────────────────────────────
# 9. finalize_job
# ────────────────────────────────────────────────────

def finalize_job(state: CVIngestionState) -> Dict:
    """Mark the ingestion job as completed."""
    job_id = state["job_id"]
    logger.info("[job=%s] Ingestion pipeline completed successfully", job_id)
    return {"status": "completed", "stage": "done"}


# ────────────────────────────────────────────────────
# 10. handle_failure
# ────────────────────────────────────────────────────

def handle_failure(state: CVIngestionState) -> Dict:
    """Handle pipeline failure — log errors and mark failed."""
    job_id = state["job_id"]
    errors = state.get("errors", [])
    logger.error("[job=%s] Pipeline FAILED with %d errors: %s", job_id, len(errors), "; ".join(errors))
    return {"status": "failed"}
